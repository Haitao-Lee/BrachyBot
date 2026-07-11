"""
Document Reader Tool
====================
Read and extract text/content from various document formats.
Supports: PDF, DOCX, TXT, CSV, JSON, MD, images (basic info)
"""

import os
import json
import csv
import logging
from typing import Dict, Any, Optional
from pathlib import Path

from tool_factory import BaseTool, ToolResult
from tool_factory.filesystem_browser import _path_is_allowed

logger = logging.getLogger(__name__)


class DocumentReaderTool(BaseTool):
    """Read and extract content from various document formats."""

    name = "doc_reader"
    description = """Read and extract content from documents.
Supported formats:
- PDF (.pdf) - Extract text from all pages
- Word (.docx) - Extract text and tables
- Text (.txt, .md, .log) - Read plain text
- CSV (.csv) - Read as table
- JSON (.json) - Read and format
- Images (.png, .jpg, .nii, .nii.gz) - Get metadata/info"""

    input_schema = {
        "file_path": {
            "type": "string",
            "description": "Path to the document file"
        },
        "action": {
            "type": "string",
            "description": "Action: read (default), summary, metadata",
            "enum": ["read", "summary", "metadata"]
        },
        "max_pages": {
            "type": "integer",
            "description": "Max pages to read for PDF (default: 10)"
        },
    }
    output_schema = {
        "success": {"type": "boolean"},
        "content": {"type": "string"},
        "metadata": {"type": "object"},
    }

    @staticmethod
    def _extractive_summary(content: str, max_chars: int = 2000) -> str:
        """Return a bounded, explicitly extractive document preview."""
        paragraphs = []
        for block in str(content or "").splitlines():
            normalized = " ".join(block.split())
            if not normalized or normalized.startswith("--- Page "):
                continue
            paragraphs.append(normalized)
            if sum(len(item) for item in paragraphs) >= max_chars:
                break
        summary = "\n".join(paragraphs)
        if len(summary) > max_chars:
            summary = summary[:max_chars].rstrip() + "..."
        return summary

    def _apply_action(self, result: ToolResult, action: str) -> ToolResult:
        if not result.success:
            return result
        data = dict(result.data or {})
        if action == "metadata":
            data["content"] = ""
            return ToolResult(
                success=True,
                data=data,
                message="Document metadata extracted",
            )
        if action == "summary":
            data["content"] = self._extractive_summary(data.get("content", ""))
            data.setdefault("metadata", {})["summary_type"] = "extractive_preview"
            return ToolResult(
                success=True,
                data=data,
                message="Extractive document summary generated",
            )
        return result

    def _read_pdf(self, file_path: str, max_pages: int = 10) -> ToolResult:
        """Read PDF file."""
        try:
            import PyPDF2
        except ImportError:
            # Try pypdf as fallback
            try:
                import pypdf
                return self._read_pdf_pypdf(file_path, max_pages)
            except ImportError:
                return ToolResult(
                    success=False,
                    error="PDF library not installed",
                    message="Please install PyPDF2: pip install PyPDF2 or pypdf"
                )

        try:
            content_parts = []
            metadata = {}

            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                # Get metadata
                if reader.metadata:
                    metadata = {
                        "title": reader.metadata.get('/Title', ''),
                        "author": reader.metadata.get('/Author', ''),
                        "pages": len(reader.pages),
                    }

                # Read pages
                pages_to_read = min(len(reader.pages), max_pages)
                for i in range(pages_to_read):
                    page = reader.pages[i]
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"--- Page {i+1} ---\n{text}")

                metadata["pages_read"] = pages_to_read
                metadata["total_pages"] = len(reader.pages)

            content = "\n\n".join(content_parts)

            # Truncate if too long
            if len(content) > 10000:
                content = content[:10000] + "\n\n... (Content truncated, total " + str(metadata.get("total_pages", 0)) + " pages)"

            return ToolResult(
                success=True,
                data={"content": content, "metadata": metadata},
                message=f"Successfully read PDF: {pages_to_read}/{metadata.get('total_pages', 0)} pages",
            )

        except Exception as e:
            logger.error(f"PDF read failed: {e}")
            return ToolResult(
                success=False,
                error=str(e),
                message=f"Failed to read PDF: {e}"
            )

    def _read_pdf_pypdf(self, file_path: str, max_pages: int = 10) -> ToolResult:
        """Read PDF using pypdf library."""
        try:
            import pypdf

            content_parts = []
            metadata = {}

            reader = pypdf.PdfReader(file_path)

            # Get metadata
            if reader.metadata:
                metadata = {
                    "title": getattr(reader.metadata, 'title', ''),
                    "author": getattr(reader.metadata, 'author', ''),
                    "pages": len(reader.pages),
                }

            # Read pages
            pages_to_read = min(len(reader.pages), max_pages)
            for i in range(pages_to_read):
                page = reader.pages[i]
                text = page.extract_text()
                if text:
                    content_parts.append(f"--- Page {i+1} ---\n{text}")

            metadata["pages_read"] = pages_to_read
            metadata["total_pages"] = len(reader.pages)

            content = "\n\n".join(content_parts)

            if len(content) > 10000:
                content = content[:10000] + "\n\n... (Content truncated)"

            return ToolResult(
                success=True,
                data={"content": content, "metadata": metadata},
                message=f"Successfully read PDF: {pages_to_read}/{metadata.get('total_pages', 0)} pages",
            )

        except Exception as e:
            return ToolResult(
                success=False,
                error=str(e),
                message=f"Failed to read PDF: {e}"
            )

    def _read_docx(self, file_path: str) -> ToolResult:
        """Read Word document."""
        try:
            from docx import Document
        except ImportError:
            return ToolResult(
                success=False,
                error="python-docx not installed",
                message="Please install python-docx: pip install python-docx"
            )

        try:
            doc = Document(file_path)

            content_parts = []
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }

            # Read paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    content_parts.append(para.text)

            # Read tables
            for table_idx, table in enumerate(doc.tables):
                content_parts.append(f"\n--- Table {table_idx + 1} ---")
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content_parts.append(row_text)

            content = "\n".join(content_parts)

            if len(content) > 10000:
                content = content[:10000] + "\n\n... (Content truncated)"

            return ToolResult(
                success=True,
                data={"content": content, "metadata": metadata},
                message=f"Successfully read Word document: {metadata['paragraphs']} paragraphs, {metadata['tables']} tables",
            )

        except Exception as e:
            return ToolResult(
                success=False,
                error=str(e),
                message=f"Failed to read Word document: {e}"
            )

    def _read_text(self, file_path: str) -> ToolResult:
        """Read plain text file."""
        try:
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                return ToolResult(
                    success=False,
                    error="Cannot decode file",
                    message="Cannot decode file. Please check the encoding format."
                )

            metadata = {
                "size_bytes": os.path.getsize(file_path),
                "lines": content.count('\n') + 1,
            }

            if len(content) > 10000:
                content = content[:10000] + "\n\n... (Content truncated)"

            return ToolResult(
                success=True,
                data={"content": content, "metadata": metadata},
                message=f"Successfully read text file: {metadata['lines']} lines",
            )

        except Exception as e:
            return ToolResult(
                success=False,
                error=str(e),
                message=f"Failed to read text file: {e}"
            )

    def _read_csv(self, file_path: str) -> ToolResult:
        """Read CSV file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)

            if not rows:
                return ToolResult(
                    success=True,
                    data={"content": "", "metadata": {"rows": 0}},
                    message="CSV file is empty"
                )

            # Format as table
            headers = rows[0]
            content_parts = ["Headers: " + " | ".join(headers)]
            content_parts.append("-" * 50)

            for i, row in enumerate(rows[1:21]):  # First 20 rows
                content_parts.append(" | ".join(row))

            if len(rows) > 21:
                content_parts.append(f"\n... ({len(rows)} rows total)")

            content = "\n".join(content_parts)
            metadata = {
                "rows": len(rows) - 1,
                "columns": len(headers),
                "headers": headers,
            }

            return ToolResult(
                success=True,
                data={"content": content, "metadata": metadata},
                message=f"Successfully read CSV: {metadata['rows']} rows, {metadata['columns']} columns",
            )

        except Exception as e:
            return ToolResult(
                success=False,
                error=str(e),
                message=f"Failed to read CSV: {e}"
            )

    def _read_json(self, file_path: str) -> ToolResult:
        """Read JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            content = json.dumps(data, indent=2, ensure_ascii=False)

            if len(content) > 10000:
                content = content[:10000] + "\n\n... (Content truncated)"

            metadata = {
                "type": type(data).__name__,
                "size": len(str(data)),
            }

            if isinstance(data, dict):
                metadata["keys"] = list(data.keys())
            elif isinstance(data, list):
                metadata["items"] = len(data)

            return ToolResult(
                success=True,
                data={"content": content, "metadata": metadata},
                message=f"Successfully read JSON: {metadata['type']}",
            )

        except Exception as e:
            return ToolResult(
                success=False,
                error=str(e),
                message=f"Failed to read JSON: {e}"
            )

    def _read_image_info(self, file_path: str) -> ToolResult:
        """Get image metadata (not content)."""
        try:
            metadata = {
                "file": os.path.basename(file_path),
                "size_bytes": os.path.getsize(file_path),
            }

            # Try to get image info
            ext = Path(file_path).suffix.lower()

            if ext in ['.nii', '.gz']:
                try:
                    import nibabel as nib
                    img = nib.load(file_path)
                    metadata["format"] = "NIfTI"
                    metadata["shape"] = list(img.shape)
                    metadata["spacing"] = list(img.header.get_zooms())
                    metadata["dtype"] = str(img.get_data_dtype())
                except Exception as exc:
                    logger.warning("Failed to read NIfTI metadata for %s: %s", file_path, exc)

            elif ext in ['.mhd', '.raw']:
                try:
                    import SimpleITK as sitk
                    img = sitk.ReadImage(file_path)
                    metadata["format"] = "MetaImage"
                    metadata["shape"] = list(img.GetSize())
                    metadata["spacing"] = list(img.GetSpacing())
                except Exception as exc:
                    logger.warning("Failed to read MetaImage metadata for %s: %s", file_path, exc)

            elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                try:
                    from PIL import Image
                    img = Image.open(file_path)
                    metadata["format"] = "Image"
                    metadata["size"] = img.size
                    metadata["mode"] = img.mode
                except Exception as exc:
                    logger.warning("Failed to read image metadata for %s: %s", file_path, exc)

            return ToolResult(
                success=True,
                data={"content": f"Image file: {metadata.get('format', 'Unknown')}", "metadata": metadata},
                message=f"Got image info: {metadata.get('format', 'Unknown')}",
            )

        except Exception as e:
            return ToolResult(
                success=False,
                error=str(e),
                message=f"Failed to get image info: {e}"
            )

    def _execute(self, **kwargs) -> ToolResult:
        file_path = kwargs.get("file_path", "")
        action = kwargs.get("action", "read")
        if action not in {"read", "summary", "metadata"}:
            return ToolResult(success=False, error="Invalid action", message="Use read, summary, or metadata")
        try:
            max_pages = max(1, min(int(kwargs.get("max_pages", 10)), 100))
        except (TypeError, ValueError):
            return ToolResult(success=False, error="Invalid max_pages", message="max_pages must be an integer")

        if not file_path:
            return ToolResult(
                success=False,
                error="No file_path provided",
                message="Please provide file path"
            )

        try:
            allowed, resolved_path = _path_is_allowed(file_path)
        except (OSError, RuntimeError, ValueError) as exc:
            return ToolResult(success=False, error=str(exc), message="Invalid document path")
        if not allowed:
            return ToolResult(
                success=False,
                error=(
                    "Access denied: path is outside the configured project/data roots. "
                    "Add it to BRACHYBOT_FILESYSTEM_ROOTS or explicitly enable trusted "
                    "global browsing with BRACHYBOT_ENABLE_FILESYSTEM_BROWSER_GLOBAL=1."
                ),
                message="Document path access restricted for security",
            )
        file_path = str(resolved_path)

        if not resolved_path.is_file():
            return ToolResult(
                success=False,
                error=f"File not found: {file_path}",
                message=f"File does not exist: {file_path}"
            )

        try:
            max_bytes = max(
                1,
                int(os.environ.get("BRACHYBOT_MAX_DOCUMENT_BYTES", str(50 * 1024 * 1024))),
            )
        except ValueError:
            max_bytes = 50 * 1024 * 1024
        if resolved_path.stat().st_size > max_bytes:
            return ToolResult(
                success=False,
                error=f"Document exceeds the configured {max_bytes}-byte limit",
                message="Document is too large to read safely",
            )

        # Get file extension
        ext = Path(file_path).suffix.lower()

        # Handle .nii.gz
        if file_path.endswith('.nii.gz'):
            ext = '.nii.gz'

        # Route to appropriate reader
        if ext == '.pdf':
            result = self._read_pdf(file_path, max_pages)
        elif ext == '.docx':
            result = self._read_docx(file_path)
        elif ext in ['.txt', '.md', '.log', '.py', '.json', '.xml', '.html']:
            if ext == '.json':
                result = self._read_json(file_path)
            else:
                result = self._read_text(file_path)
        elif ext == '.csv':
            result = self._read_csv(file_path)
        elif ext in ['.nii', '.nii.gz', '.mhd', '.raw', '.dcm',
                      '.png', '.jpg', '.jpeg', '.bmp', '.gif']:
            result = self._read_image_info(file_path)
        else:
            # Try as text
            result = self._read_text(file_path)
        return self._apply_action(result, action)
